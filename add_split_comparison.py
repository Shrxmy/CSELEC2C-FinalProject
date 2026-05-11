from pathlib import Path
import csv
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC = Path('paper/GTD_Attack_Success_Manuscript_CART.docx')
TEMP_METRICS = Path('results/tables/model_metrics.csv')
RAND_METRICS = Path('results/tables/random_split/model_metrics.csv')
OUT_TABLE = Path('results/tables/temporal_vs_random_split_comparison.csv')
OUT_FIG = Path('results/figures/temporal_vs_random_split_comparison.png')

name_map = {
    'logistic_regression': 'Logistic Regression',
    'cart_decision_tree': 'CART',
    'xgboost': 'XGBoost',
    'feedforward_torch_nn': 'Feedforward NN',
}
order = ['logistic_regression', 'cart_decision_tree', 'xgboost', 'feedforward_torch_nn']

# Build comparison data.
temporal = pd.read_csv(TEMP_METRICS)
random = pd.read_csv(RAND_METRICS)
t = temporal[temporal['split'] == 'test'].set_index('model')
r = random[random['split'] == 'test'].set_index('model')
rows = []
for model in order:
    rows.append({
        'model': model,
        'model_label': name_map[model],
        'temporal_roc_auc': t.loc[model, 'roc_auc'],
        'random_roc_auc': r.loc[model, 'roc_auc'],
        'delta_roc_auc': r.loc[model, 'roc_auc'] - t.loc[model, 'roc_auc'],
        'temporal_f1': t.loc[model, 'f1'],
        'random_f1': r.loc[model, 'f1'],
        'delta_f1': r.loc[model, 'f1'] - t.loc[model, 'f1'],
        'temporal_balanced_accuracy': t.loc[model, 'balanced_accuracy'],
        'random_balanced_accuracy': r.loc[model, 'balanced_accuracy'],
        'delta_balanced_accuracy': r.loc[model, 'balanced_accuracy'] - t.loc[model, 'balanced_accuracy'],
        'temporal_mcc': t.loc[model, 'mcc'],
        'random_mcc': r.loc[model, 'mcc'],
        'delta_mcc': r.loc[model, 'mcc'] - t.loc[model, 'mcc'],
    })
comparison = pd.DataFrame(rows)
OUT_TABLE.parent.mkdir(parents=True, exist_ok=True)
comparison.to_csv(OUT_TABLE, index=False)

# Create comparison figure.
plot_df = comparison.melt(
    id_vars=['model_label'],
    value_vars=['temporal_roc_auc', 'random_roc_auc', 'temporal_f1', 'random_f1', 'temporal_balanced_accuracy', 'random_balanced_accuracy'],
    var_name='metric_split',
    value_name='score',
)
plot_df['split'] = plot_df['metric_split'].str.extract(r'^(temporal|random)')
plot_df['metric'] = plot_df['metric_split'].str.replace('temporal_', '', regex=False).str.replace('random_', '', regex=False)
plot_df['metric'] = plot_df['metric'].map({'roc_auc': 'ROC-AUC', 'f1': 'F1', 'balanced_accuracy': 'Balanced Accuracy'})
sns.set_theme(style='whitegrid')
g = sns.catplot(
    data=plot_df,
    x='model_label', y='score', hue='split', col='metric',
    kind='bar', height=4, aspect=1.05, sharey=True,
)
g.set_axis_labels('', 'Score')
g.set_titles('{col_name}')
for ax in g.axes.flat:
    ax.tick_params(axis='x', rotation=25)
    ax.set_ylim(0, 1)
g.fig.suptitle('Temporal vs Random Split Test Performance', y=1.08)
g.tight_layout()
OUT_FIG.parent.mkdir(parents=True, exist_ok=True)
g.savefig(OUT_FIG, dpi=180, bbox_inches='tight')
plt.close('all')

# Word helpers.
def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def format_body(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)


def format_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(8)
        run.italic = True


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(6)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in('w:tcBorders')
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                element = tcBorders.find(qn('w:' + edge))
                if element is None:
                    element = OxmlElement('w:' + edge)
                    tcBorders.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '808080')


def remove_existing_split_block(doc):
    start_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Temporal and Random Split Comparison':
            start_idx = i
        elif start_idx is not None and p.text.strip().startswith('Figure 1.'):
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        for p in list(doc.paragraphs[start_idx:end_idx]):
            p._element.getparent().remove(p._element)
    # Remove old comparison table if rerun.
    for table in list(doc.tables):
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == 'Model' and len(table.rows[0].cells) >= 9:
            headers = [c.text.strip() for c in table.rows[0].cells]
            if 'Temp ROC' in headers or 'Rand ROC' in headers:
                table._element.getparent().remove(table._element)


doc = Document(DOC)
remove_existing_split_block(doc)

# Insert after Table 1 caption, before Figure 1.
anchor = next((p for p in doc.paragraphs if p.text.strip().startswith('Table 1. Test-set statistical summary')), None)
if anchor is None:
    raise RuntimeError('Could not find Table 1 caption anchor')

head = insert_paragraph_after(anchor, 'Temporal and Random Split Comparison', style='Heading 1')
head.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro = insert_paragraph_after(head, 'We also ran the same four-model pipeline with a stratified random split to check how much the evaluation design affected performance. The random split uses 70% of the data for training, 15% for validation, and 15% for testing. As expected, the random split produced higher scores because events from similar periods can appear in all partitions. The temporal split remains the stricter and more realistic result for the manuscript, while the random split is useful as a comparison check.')
format_body(intro)

headers = ['Model', 'Temp ROC', 'Rand ROC', 'Delta ROC', 'Temp F1', 'Rand F1', 'Delta F1', 'Temp BalAcc', 'Rand BalAcc']
table = doc.add_table(rows=1 + len(comparison), cols=len(headers))
for j, h in enumerate(headers):
    table.rows[0].cells[j].text = h
for i, row in comparison.iterrows():
    vals = [
        row['model_label'], row['temporal_roc_auc'], row['random_roc_auc'], row['delta_roc_auc'],
        row['temporal_f1'], row['random_f1'], row['delta_f1'], row['temporal_balanced_accuracy'], row['random_balanced_accuracy'],
    ]
    for j, v in enumerate(vals):
        table.rows[i+1].cells[j].text = v if j == 0 else f'{float(v):.3f}'
style_table(table)
intro._p.addnext(table._tbl)
cap = insert_paragraph_after(intro, 'Table 2. Temporal-vs-random split test comparison. Delta values show random split minus temporal split.')
table._tbl.addnext(cap._p)
format_caption(cap)

pimg = insert_paragraph_after(cap)
pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
pimg.add_run().add_picture(str(OUT_FIG), width=Inches(3.15))
cap2 = insert_paragraph_after(pimg, 'Figure 1. Temporal and random split comparison for ROC-AUC, F1, and balanced accuracy.')
format_caption(cap2)

note = insert_paragraph_after(cap2, 'The comparison confirms that the random split is easier: every model improved in ROC-AUC, F1, balanced accuracy, and MCC. XGBoost and the feedforward neural network still remained the strongest models overall. However, because random splitting mixes years across train and test data, we treat it as a robustness check rather than the main evidence.')
format_body(note)

# Renumber later figure captions by +1 because this inserts a new Figure 1 before the old Figure 1.
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('Figure '):
        # skip the new split figure caption exactly
        if text.startswith('Figure 1. Temporal and random split comparison'):
            continue
        try:
            num = int(text.split('.')[0].split()[1])
        except Exception:
            continue
        if num >= 1:
            # preserve style roughly
            new_text = text.replace(f'Figure {num}.', f'Figure {num+1}.', 1)
            p.clear()
            p.add_run(new_text)
            format_caption(p)

# Update earlier text that references figure list if needed.
for p in doc.paragraphs:
    if p.text.startswith('The notebook also produced ROC curves'):
        p.clear()
        p.add_run('The notebook also produced ROC curves, precision-recall curves, calibration plots, split-comparison plots, overfitting checks, generalization-gap plots, learning curves, and feature-importance graphics. These visualizations are saved in the results/figures directory. Each plot emphasizes something different: ROC curves show ranking, precision-recall curves show behavior under imbalance, split-comparison plots show how evaluation design changes performance, overfitting graphs show train-validation-test stability, and feature-importance plots show which variables carry much of the signal.')
        format_body(p)
        break

try:
    doc.save(DOC)
    saved_to = DOC
except PermissionError:
    saved_to = DOC.with_name(DOC.stem + '_split_comparison.docx')
    doc.save(saved_to)
print(f'Wrote {OUT_TABLE}')
print(f'Wrote {OUT_FIG}')
print(f'Updated document: {saved_to}')
